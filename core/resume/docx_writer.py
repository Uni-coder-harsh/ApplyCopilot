"""
DOCX writer — generates a clean, professional resume in .docx format.
Uses python-docx with consistent styles. ATS-friendly layout.
No tables, no text boxes — plain sections that parse correctly.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.resume.builder import ResumeData

logger = logging.getLogger(__name__)

# Colours
COLOR_NAME = RGBColor(0x1A, 0x1A, 0x2E)      # dark navy
COLOR_SECTION = RGBColor(0x16, 0x21, 0x3E)    # section headers
COLOR_ACCENT = RGBColor(0x0F, 0x3D, 0x85)     # links and highlights
COLOR_BODY = RGBColor(0x2C, 0x2C, 0x2C)       # body text
COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)      # secondary text


def generate_docx(
    resume_data: ResumeData,
    output_path: Path,
) -> Path:
    """
    Generate a DOCX resume from ResumeData and save to output_path.
    Returns the path to the saved file.
    """
    doc = Document()
    _set_margins(doc)
    _set_default_style(doc)

    # Header: name + contact
    _write_header(doc, resume_data)

    # Summary
    summary = resume_data.tailored_summary or resume_data.summary
    if summary:
        _write_section(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(summary)
        _style_body(p)

    # Skills
    if resume_data.skill_groups:
        _write_section(doc, "SKILLS")
        for group in resume_data.skill_groups:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{group.category}: ")
            run_label.bold = True
            run_label.font.size = Pt(10)
            run_label.font.color.rgb = COLOR_BODY
            run_value = p.add_run(", ".join(group.skills))
            run_value.font.size = Pt(10)
            run_value.font.color.rgb = COLOR_BODY
            p.paragraph_format.space_after = Pt(2)

    # Projects
    projects = resume_data.tailored_projects or resume_data.projects
    if projects:
        _write_section(doc, "PROJECTS")
        for project in projects[:5]:
            _write_project(doc, project)

    # Education
    if resume_data.education:
        _write_section(doc, "EDUCATION")
        for edu in resume_data.education:
            _write_education(doc, edu)

    doc.save(str(output_path))
    logger.info(f"Resume saved to {output_path}")
    return output_path


def _set_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def _set_default_style(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = COLOR_BODY


def _write_header(doc: Document, resume_data: ResumeData):
    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume_data.full_name.upper())
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = COLOR_NAME
    name_para.paragraph_format.space_after = Pt(4)

    # Contact line
    contact_parts = []
    if resume_data.email:
        contact_parts.append(resume_data.email)
    if resume_data.phone:
        contact_parts.append(resume_data.phone)
    if resume_data.location:
        contact_parts.append(resume_data.location)
    if resume_data.linkedin:
        contact_parts.append(resume_data.linkedin)
    if resume_data.github:
        contact_parts.append(resume_data.github)
    if resume_data.portfolio:
        contact_parts.append(resume_data.portfolio)

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_MUTED
        contact_para.paragraph_format.space_after = Pt(6)

    _add_horizontal_rule(doc)


def _write_section(doc: Document, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLOR_SECTION
    _add_bottom_border(p)


def _write_project(doc: Document, project):
    # Project name + dates on same line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)

    name_run = p.add_run(project.name)
    name_run.bold = True
    name_run.font.size = Pt(10.5)
    name_run.font.color.rgb = COLOR_BODY

    # Date range
    dates = []
    if project.start_date:
        dates.append(project.start_date)
    if project.end_date:
        dates.append(project.end_date)
    if dates:
        date_run = p.add_run(f"  |  {' – '.join(dates)}")
        date_run.font.size = Pt(9.5)
        date_run.font.color.rgb = COLOR_MUTED

    # Tech stack
    if project.tech_stack:
        tech_p = doc.add_paragraph()
        tech_p.paragraph_format.space_after = Pt(1)
        tech_run = tech_p.add_run("Tech: " + ", ".join(project.tech_stack))
        tech_run.font.size = Pt(9.5)
        tech_run.font.color.rgb = COLOR_MUTED
        tech_run.italic = True

    # Description
    if project.description:
        desc_p = doc.add_paragraph(project.description)
        _style_body(desc_p)

    # Impact
    if project.impact:
        impact_p = doc.add_paragraph()
        impact_run = impact_p.add_run("Impact: " + project.impact)
        impact_run.font.size = Pt(10)
        impact_run.font.color.rgb = COLOR_BODY

    # Repo link
    if project.repo_url:
        repo_p = doc.add_paragraph()
        repo_run = repo_p.add_run(project.repo_url)
        repo_run.font.size = Pt(9.5)
        repo_run.font.color.rgb = COLOR_ACCENT


def _write_education(doc: Document, edu):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)

    inst_run = p.add_run(edu.institution)
    inst_run.bold = True
    inst_run.font.size = Pt(10.5)

    degree_parts = [edu.degree, edu.field]
    degree_str = " in ".join(filter(None, degree_parts))
    if degree_str:
        deg_run = p.add_run(f"  —  {degree_str}")
        deg_run.font.size = Pt(10)
        deg_run.font.color.rgb = COLOR_MUTED

    # Year + CGPA
    meta_parts = []
    years = [str(y) for y in [edu.start_year, edu.end_year] if y]
    if years:
        meta_parts.append(" – ".join(years))
    if edu.cgpa:
        meta_parts.append(f"CGPA: {edu.cgpa}")

    if meta_parts:
        meta_p = doc.add_paragraph("  ".join(meta_parts))
        for run in meta_p.runs:
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_MUTED
        meta_p.paragraph_format.space_after = Pt(2)


def _style_body(para):
    para.paragraph_format.space_after = Pt(3)
    for run in para.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_BODY


def _add_horizontal_rule(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "162139")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "162139")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_output_filename(
    resume_data: ResumeData,
    version: Optional[str] = None,
) -> str:
    """Generate a clean output filename."""
    name = resume_data.full_name.replace(" ", "_").lower()
    company = (resume_data.job_company or "general").replace(" ", "_").lower()
    role = (resume_data.job_role or "resume").replace(" ", "_").lower()
    date_str = datetime.now().strftime("%Y%m%d")
    version_str = f"_v{version}" if version else ""
    return f"{name}_{company}_{role}{version_str}_{date_str}.docx"
